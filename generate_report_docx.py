import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_title_page(doc):
    doc.add_paragraph('\n'*5)
    title = doc.add_heading('NIST University\nAutomated Assignment Evaluation System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n'*2)
    subtitle = doc.add_paragraph('Comprehensive Project Report')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(20)
    doc.add_paragraph('\n'*5)
    info = doc.add_paragraph('Prepared for: Appathon Submission\nDeveloped by: NIST Student Developer Team\nDate: May 2026')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

def add_chapter(doc, title, content_paragraphs, include_break=True):
    doc.add_heading(title, level=1)
    for p in content_paragraphs:
        if isinstance(p, tuple) and p[0] == 'sub':
            doc.add_heading(p[1], level=2)
        elif isinstance(p, tuple) and p[0] == 'bullet':
            for item in p[1]:
                doc.add_paragraph(item, style='List Bullet')
        else:
            doc.add_paragraph(p)
    if include_break:
        doc.add_page_break()

def main():
    doc = Document()
    
    # Title Page
    add_title_page(doc)
    
    # Abstract
    add_chapter(doc, 'Abstract', [
        "The Automated Assignment Evaluation System is a comprehensive web-based platform designed for NIST University to streamline the grading process for both faculty and students. In modern educational environments, manual grading of descriptive essays and coding assignments is time-consuming and prone to human error or bias.",
        "This system leverages Natural Language Processing (NLP) and sandboxed code execution environments to automatically evaluate student submissions against dynamic, faculty-defined rubrics. Furthermore, it integrates a robust plagiarism detection engine utilizing TF-IDF and Cosine Similarity to ensure academic integrity.",
        "The platform offers distinct dashboards for faculty to manage assignments and batch-evaluate submissions, and for students to track their performance and view detailed feedback reports. The UI is completely branded with NIST University's official color palette and typography, providing a premium, native experience for the university ecosystem.",
        "By automating the tedious aspects of grading, the system empowers educators to focus on curriculum development and personalized student mentorship, while providing students with instant, objective feedback."
    ])

    # Chapter 1: Introduction
    add_chapter(doc, '1. Introduction', [
        ('sub', '1.1 Background'),
        "Educational institutions worldwide face a growing challenge in efficiently evaluating student assignments. As class sizes increase, the burden on faculty to provide timely and detailed feedback grows exponentially. This often leads to delayed results and generic feedback, which hampers the learning experience.",
        ('sub', '1.2 Problem Statement'),
        "Currently, NIST University relies on manual or semi-automated processes for assignment collection and grading. These processes are fragmented. Code execution is done manually by downloading zip files, running them locally, and comparing outputs. Descriptive answers are read individually without systematic rubrics. Plagiarism checking requires external, often paid, tools.",
        ('sub', '1.3 Objectives'),
        ('bullet', [
            "Develop a unified platform for assignment distribution, submission, and grading.",
            "Automate the evaluation of coding assignments across multiple languages (Python, Java, C++).",
            "Automate the grading of descriptive text assignments using Natural Language Processing.",
            "Implement an internal peer-to-peer plagiarism detection system.",
            "Provide role-based dashboards with analytics for both faculty and students.",
            "Ensure the UI conforms perfectly to NIST University branding."
        ]),
        ('sub', '1.4 Scope'),
        "The scope of this project encompasses the entire lifecycle of an assignment, from creation by a faculty member, defining dynamic rubrics and test cases, student submission (text, PDF, or code files), automated evaluation, to final grade publishing and analytics reporting."
    ])

    # Chapter 2: System Architecture
    add_chapter(doc, '2. System Architecture & Tech Stack', [
        "The system follows a monolithic Model-View-Controller (MVC) architecture, optimized for rapid deployment and ease of maintenance within a university server environment.",
        ('sub', '2.1 Technology Stack'),
        ('bullet', [
            "Backend Framework: Python with Flask.",
            "Database: SQLite (via SQLAlchemy ORM) for lightweight, transactional data storage.",
            "Frontend: HTML5, Vanilla CSS with custom CSS variables for theming, and Vanilla JavaScript.",
            "NLP Processing: TextBlob and NLTK for sentiment, spelling, and semantic analysis.",
            "Code Execution: Python's built-in subprocess module with strict timeouts.",
            "Authentication: Flask-Login for secure session management and role-based access control."
        ]),
        ('sub', '2.2 Application Flow'),
        "1. Faculty creates an assignment, defining the type (coding or descriptive), reference solution, rubric weights, and (if coding) input/output test cases.",
        "2. Students log in, view pending assignments, and upload their work (PDF, TXT, PY, CPP, JAVA).",
        "3. The system parses the uploaded file, extracting raw text.",
        "4. Faculty clicks 'Batch Evaluate'. The system queues submissions.",
        "5. The Code Runner or NLP Evaluator processes the text. The Plagiarism module compares it against peers.",
        "6. A final score is aggregated based on rubric weights and stored in the database.",
        "7. Faculty reviews and releases grades. Students view their detailed feedback reports."
    ])

    # Chapter 3: Database Models
    add_chapter(doc, '3. Database Design', [
        "The database is normalized and utilizes SQLAlchemy for object-relational mapping. The core entities are:",
        ('sub', '3.1 User Model'),
        "Stores authentication details, role ('faculty' or 'student'), full name, and university-issued student IDs. This model enforces role-based access control across the application.",
        ('sub', '3.2 Assignment Model'),
        "Holds the assignment metadata (title, description, due date), configuration (type, language), and the reference solution. It also stores JSON-encoded test cases for coding tasks.",
        ('sub', '3.3 Rubric Model'),
        "Linked to an Assignment via a one-to-many relationship. It defines the criteria (e.g., 'Correctness', 'Code Quality', 'Originality') and their respective percentage weights totaling 100%.",
        ('sub', '3.4 Submission Model'),
        "Records the student's submission, file path, extracted raw text, and current processing status ('pending', 'evaluated', 'failed').",
        ('sub', '3.5 Evaluation Model'),
        "Stores the final results, including the aggregated score, letter grade, granular sub-scores (correctness, style, originality), detailed JSON feedback, test case execution logs, and plagiarism match data."
    ])

    # Chapter 4: Core Evaluation Modules
    add_chapter(doc, '4. Core Evaluation Modules', [
        "The intelligence of the platform lies in its modular evaluation services. By decoupling these services, the system can easily be extended to support new assignment types in the future.",
        ('sub', '4.1 File Parsing Engine'),
        "Handles the extraction of raw text from uploaded files. Supported formats include .txt, source code extensions (.py, .java, .cpp), and PDFs (via PyMuPDF). It features robust error handling to gracefully degrade if a file is corrupted, returning empty strings to prevent pipeline crashes.",
        ('sub', '4.2 Code Execution Sandbox'),
        "Executes untrusted student code in an isolated environment. It writes the code to a temporary directory, compiles it if necessary (g++, javac), and feeds input test cases via stdin. The standard output is captured and strictly compared against the expected output. Strict timeouts (e.g., 5 seconds) are enforced to prevent infinite loops from hanging the server.",
        ('sub', '4.3 NLP Descriptive Evaluator'),
        "Evaluates essays and reports against a reference solution provided by the faculty. It uses cosine similarity to measure semantic overlap. It also analyzes text for spelling errors, grammar consistency, clarity, and completeness, generating constructive feedback bullet points automatically.",
        ('sub', '4.4 Plagiarism Detection'),
        "A highly optimized peer-to-peer comparison engine. It uses n-gram overlap and TF-IDF cosine similarity to detect both exact copying and heavy paraphrasing. To ensure system stability, large texts are truncated, and the mathematical formula strictly normalizes results between 0% and 100%."
    ])

    # Chapter 5: Security & Exception Handling
    add_chapter(doc, '5. Security & Exception Handling', [
        "Given the nature of executing user-provided code and processing file uploads, security and stability are paramount.",
        ('sub', '5.1 Input Validation & Limits'),
        ('bullet', [
            "File Size Limits: A strict 16MB cap is enforced at the Flask application layer (MAX_CONTENT_LENGTH) to prevent Denial of Service (DoS) attacks via disk exhaustion.",
            "Rubric Validation: The system enforces that rubric weights must sum to a positive integer and dynamically normalizes them during evaluation if they do not exactly equal 100%.",
            "Type Checking: Uploaded files are verified against a whitelist of allowed extensions."
        ]),
        ('sub', '5.2 Execution Security'),
        "The subprocess environment uses timeouts and captures stderr separately. Missing dependencies (like missing compilers on the host system) are caught cleanly, and informative error messages are returned to the user rather than generic 500 Internal Server Errors.",
        ('sub', '5.3 Access Control'),
        "Strict Route Decorators ensure that students cannot access faculty dashboards or evaluation triggers. Furthermore, ownership checks are implemented so students can only view their own evaluation reports, protecting peer privacy."
    ])

    # Chapter 6: Batch Evaluation System
    add_chapter(doc, '6. Batch Evaluation & Faculty Workflow', [
        "To accommodate classes of hundreds of students, the platform implements a Batch Evaluation workflow.",
        ('sub', '6.1 One-Click Processing'),
        "Faculty can trigger the evaluation of all 'pending' submissions simultaneously via the 'Batch Evaluate' button. The system iterates through the queue, applying the relevant parsers, runners, and plagiarism checks.",
        ('sub', '6.2 Grade Withholding (Release Toggles)'),
        "To prevent students from seeing grades before the faculty has reviewed them, evaluations are generated in an 'unreleased' state. Faculty can review the auto-generated feedback, make manual adjustments if necessary, and then click 'Release Grades' to make them visible on the student portal.",
        ('sub', '6.3 HTML Export'),
        "For official university record-keeping, any evaluation report can be exported as a standalone, styled HTML document."
    ])

    # Chapter 7: Testing & Validation
    add_chapter(doc, '7. Testing & Validation Results', [
        "Extensive testing was conducted to ensure the reliability of the Automated Evaluator.",
        ('sub', '7.1 Mock Data Seeding'),
        "A rigorous seed script was developed to generate 21 diverse submissions across 15 mock student accounts. This dataset included perfect solutions, syntax errors, logical errors (e.g., linear search instead of binary search), and highly plagiarized essays.",
        ('sub', '7.2 Test Outcomes'),
        ('bullet', [
            "Accuracy: The code runner correctly failed logical errors and assigned 0 Correctness points, while penalizing syntax errors immediately.",
            "Plagiarism Mathematics: Initial mathematical inflation bugs in the cosine similarity formula were identified and patched, ensuring originality scores always calculate correctly between 0% and 100%.",
            "Performance: The introduction of a 50,000 character truncation limit resolved hanging issues during massive n-gram comparisons."
        ])
    ])

    # Chapter 8: Conclusion
    add_chapter(doc, '8. Conclusion & Future Enhancements', [
        "The NIST University Automated Assignment Evaluation System successfully demonstrates a robust, scalable approach to modernizing educational workflows. By combining sandboxed code execution, NLP-driven descriptive grading, and peer plagiarism detection into a single, branded platform, it significantly reduces administrative overhead.",
        ('sub', '8.1 Future Scope'),
        ('bullet', [
            "Containerization: Migrating the code runner from local subprocesses to ephemeral Docker containers for enhanced security and scalability.",
            "LLM Integration: Replacing basic NLP heuristics with locally hosted Large Language Models (LLMs) via Ollama for more nuanced, semantic feedback on essays.",
            "Asynchronous Task Queues: Implementing Celery and Redis to handle batch evaluations asynchronously, preventing HTTP timeouts for extremely large classes.",
            "LMS Integration: Adding LTI or REST API hooks to sync grades directly back to the university's main Learning Management System."
        ])
    ], include_break=False)

    doc.save('NIST_AutoEval_Project_Report.docx')
    print("Report generated successfully as NIST_AutoEval_Project_Report.docx")

if __name__ == '__main__':
    main()
